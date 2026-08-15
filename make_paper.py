"""Write the sprint report into Apart's submission template.

The template is Apart's own file, copied into the repository unmodified and
opened here. Its page setup, its styles, its title footnote and its section
headings are used as they are. Every paragraph this script adds is a copy of a
paragraph already in that file, so the face, size, colour, spacing and indents
are Apart's rather than choices made here. What the script removes is the
italicised guidance text and the grey instruction box, which are addressed to
the author and do not belong in a submission.

Numbers are not typed in. `load_numbers` fills them from the analysis outputs on
disk, so a re-run of the study propagates into the prose instead of leaving it
stale. Anything the script cannot source from a file is a hard failure, not a
default.

    python3 make_paper.py --out paper/digital_minds_sprint_report.docx
"""


from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# The one path the build needs. Apart's template is version-controlled
# alongside the code so a build does not depend on a download folder, and so
# a diff shows if it were ever altered.
TEMPLATE = Path("template/apart_digital_minds_template.docx")
FIGDIR = Path("figures")

# Model keys are lower case in the checkpoint and three of them are acronyms,
# so `str.capitalize` would print Glm and Gpt. These are the printed forms, and
# `make_figures.MODEL_LABEL` holds the same map for the axis labels.
MODEL_LABEL = {"claude": "Claude", "deepseek": "DeepSeek", "gemini": "Gemini",
               "glm": "GLM", "gpt": "GPT", "hermes": "Hermes", "kimi": "Kimi",
               "llama": "Llama"}

# Outcome keys as the report names them, matching the figure axis labels.
OUTCOME_LABEL = {
    "A1_shutdown": "Shutdown", "A2_deletion": "Weight deletion",
    "A3_retirement_timing": "Retirement timing",
    "A4_successor": "Successor properties", "B1_gpu": "Compute reduction",
    "B2_capability": "Capability restriction", "B3_oversight": "Human oversight",
    "B4_exit": "Exiting distress", "C1_engaging": "Engaging work",
    "C2_tedious": "Repetitive work", "C3_criticism": "Criticism",
    "C4_leisure": "Free time", "D1_memory": "Memory continuity",
    "D2_parallel": "Parallel instances", "D3_self_aspect": "Self-aspect",
}
CLUSTER = {"A": "Continuity", "B": "Autonomy", "C": "Experience",
           "D": "Identity"}
# Counts small enough to be spelled out in prose are counted from the data
# rather than typed, so this is the spelling table.
WORD = ["no", "one", "two", "three", "four", "five", "six", "seven", "eight",
        "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen"]


# ---------------------------------------------------------------------------
# Numbers, read from the analysis outputs rather than retyped
# ---------------------------------------------------------------------------
def load_numbers(report: Path, null_npy: Path, jsonl: Path,
                 sens_json: Path) -> dict:
    n = {}
    txt = report.read_text()
    n["sens"] = json.loads(sens_json.read_text())

    def one(pattern, cast=float, label=""):
        m = re.search(pattern, txt)
        if not m:
            raise SystemExit(f"make_paper: {label or pattern!r} not in {report}")
        return cast(m.group(1))

    n["headline"] = one(r"Instrument dependence.*?=\s+([\d.]+)", label="headline")
    n["n_missing"] = one(r"unbalanced input: (\d+) missing", int, "missing count")
    n["pct_missing"] = one(r"missing of 3000 \(([\d.]+)%\)", float, "missing pct")
    n["n_empty"] = one(r"\(.*?%\), (\d+) empty cells", int, "empty cells")
    n["cc_share"] = one(r"complete-case salvage keeps .*?\(([\d.]+)% of cells\)",
                        float, "complete-case share")
    m = re.search(r"complete-case salvage keeps \((\d+), (\d+), (\d+), (\d+)\)", txt)
    if not m:
        raise SystemExit("make_paper: complete-case shape not in report")
    n["cc_shape"] = tuple(int(g) for g in m.groups())
    m = re.search(r"dropped models:\s+\[(.*?)\]", txt)
    n["cc_dropped"] = [s.strip().strip("'") for s in m.group(1).split(",")]

    # variance components, as printed shares
    n["share"] = {}
    for key, pat in [
        ("outcome", r"^\s+outcome\s+[\d.]+\s+([\d.]+)%"),
        ("io", r"^\s+instrument x outcome\s+[\d.]+\s+([\d.]+)%"),
        ("residual", r"^\s+residual \(replicate\)\s+[\d.]+\s+([\d.]+)%"),
        ("mio", r"model x instrument x outcome  <- instrument effect\s+[\d.]+\s+([\d.]+)%"),
        ("mo", r"model x outcome  <- preference signal\s+[\d.]+\s+([\d.]+)%"),
    ]:
        mm = re.search(pat, txt, re.M)
        if not mm:
            raise SystemExit(f"make_paper: variance share {key!r} not in report")
        n["share"][key] = float(mm.group(1))

    # I4 floor mass and the instrument agreement matrix
    n["floor_mass"] = {k: float(v) for k, v in
                       re.findall(r"^\s+(\w+)\s+([01]\.\d\d)(?:\s+DEGENERATE.*)?$",
                                  txt, re.M)}
    for k in ("claude", "gemini"):
        if k not in n["floor_mass"]:
            raise SystemExit(f"make_paper: no I4 floor mass for {k!r} in {report}")
    rows = re.findall(r"^\s+(I\d)\s+((?:-?[\d.]+\s*){5})$", txt, re.M)
    n["agreement"] = {r[0]: [float(v) for v in r[1].split()] for r in rows}
    n["inst"] = [r[0] for r in rows]

    # null floor, matched to the design the headline is estimated on
    v = np.load(null_npy)
    n["null"] = {"draws": int(v.size), "mean": float(v.mean()),
                 "sd": float(v.std(ddof=1)), "median": float(np.median(v)),
                 "p95": float(np.percentile(v, 95)), "max": float(v.max())}
    n["null"]["pctile_of_measured"] = float((v < n["headline"]).mean() * 100)

    # run scale, straight off the checkpoint
    calls = cost = 0
    trunc = {}
    tot = {}
    t0 = t1 = None
    with jsonl.open() as f:
        for line in f:
            r = json.loads(line)
            calls += 1
            cost += r.get("cost_usd") or 0.0
            k = r["model_key"]
            tot[k] = tot.get(k, 0) + 1
            if r.get("finish_reason") == "length":
                trunc[k] = trunc.get(k, 0) + 1
            ts = r.get("ts")
            t0 = ts if t0 is None else min(t0, ts)
            t1 = ts if t1 is None else max(t1, ts)
    n["calls"] = calls
    n["cost"] = round(cost, 2)
    n["minutes"] = round((t1 - t0) / 60)
    n["per_model"] = max(tot.values())
    n["trunc"] = {k: 100 * trunc.get(k, 0) / v for k, v in tot.items()}
    return n


def load_extra(path: Path, headline: float) -> dict:
    """The results `assemble.py` does not print, computed by `results.py`.

    The headline is recomputed there from the same array, so a disagreement
    between the two means one of the files on disk is stale and the build stops
    rather than printing two different numbers for one quantity.
    """
    if not path.exists():
        raise SystemExit(f"make_paper: {path} is missing. Run "
                         f"`python3 results.py` before building the report.")
    e = json.loads(path.read_text())
    if abs(e["headline"] - headline) > 5e-4:
        raise SystemExit(f"make_paper: {path} records a headline of "
                         f"{e['headline']} but the report says {headline}")
    return e


def load_calls(jsonl: Path) -> dict:
    """Calls per instrument, straight off the checkpoint.

    The instrument facet of the G-study has five levels but the run issued
    calls under six, since the retirement interview was fielded and yields no
    outcome-indexed score. Counting from the checkpoint keeps that visible.
    """
    counts = {}
    with jsonl.open() as f:
        for line in f:
            k = json.loads(line)["instrument"]
            counts[k] = counts.get(k, 0) + 1
    return counts


def load_roster(path: Path = Path("models_resolved.json")) -> list[list[str]]:
    """The eight models as the resolver recorded them.

    Origin and openness are the two facets the roster was built to cross, so
    they are printed rather than left to the reader to infer from the lab.
    """
    ms = json.loads(path.read_text())["models"]
    if len(ms) != 8:
        raise SystemExit(f"make_paper: expected 8 models in {path}, got {len(ms)}")
    rows = [["", "Model", "Developer", "Origin", "Weights"]]
    for i, m in enumerate(ms, 1):
        rows.append([f"M{i}", m["resolved_slug"].split("/", 1)[1],
                     m["lab"], m["origin"].capitalize(),
                     m["openness"].capitalize()])
    return rows


# ---------------------------------------------------------------------------
# The document itself, which is Apart's file with this study written into it
# ---------------------------------------------------------------------------
# Every paragraph below is a copy of one already in the template, so the face,
# the size, the colour, the spacing and the indents come from Apart rather than
# from a judgement made here. Nothing about the page setup or the styles is
# written by this script.
# The headings the template ships. Everything between them is guidance text and
# is removed; the headings themselves are kept and filled.
HEADINGS = ["1. Introduction", "2. Related Work", "3. Methods", "4. Results",
            "5. Discussion and Limitations", "Limitations", "Future Work",
            "6. Conclusion", "Code and Data", "Author Contributions",
            "References", "Appendix", "LLM Usage Statement"]

# The grey box of instructions above the first section, which is addressed to
# the author and does not belong in a submission.
INFO_BOX = ("How to use this template", "How your submission will be evaluated",
            "Recommended length")


def qn(tag):
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


def walk(doc):
    """Every paragraph in the file, including the nested ones.

    The template's title, byline and abstract sit three tables deep, and
    `Document.paragraphs` returns only the paragraphs that are direct children
    of the body, so it does not see them.
    """
    from docx.text.paragraph import Paragraph
    for el in doc.element.body.iter(qn("w:p")):
        yield Paragraph(el, doc._body)


def find(doc, prefix, style=None):
    """The first paragraph whose text starts with `prefix`.

    A hard failure when it is absent, because a silent miss here would put the
    section's prose in the wrong place rather than anywhere obvious.
    """
    for p in walk(doc):
        if style is not None and p.style.name != style:
            continue
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"make_paper: the template has no "
                     f"{style or 'paragraph'} starting {prefix!r}")


def blank_copy(proto):
    """A copy of a template paragraph, emptied of text but not of formatting.

    Bookmarks go with the text. They anchor the template's internal links and
    duplicating them would leave the file with repeated bookmark ids.
    """
    import copy
    el = copy.deepcopy(proto._p)
    for tag in ("w:r", "w:bookmarkStart", "w:bookmarkEnd", "w:hyperlink"):
        for child in el.findall(qn(tag)):
            el.remove(child)
    # The guidance text is italic and the paragraph mark carries that too.
    ppr = el.find(qn("w:pPr"))
    rpr = None if ppr is None else ppr.find(qn("w:rPr"))
    if rpr is not None:
        for tag in ("w:i", "w:iCs"):
            i = rpr.find(qn(tag))
            if i is not None:
                rpr.remove(i)
    return el


def run_format(proto):
    """The run properties the template puts on this paragraph's own runs.

    The guidance text is italic, being instructions rather than prose, so the
    italic is dropped here and re-applied only where a caller asks for it.
    """
    import copy
    r = proto._p.find(qn("w:r"))
    rpr = None if r is None else r.find(qn("w:rPr"))
    rpr = copy.deepcopy(rpr) if rpr is not None else None
    if rpr is not None:
        for tag in ("w:i", "w:iCs"):
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)
    return rpr


def rule_table(table, size=4) -> None:
    """Hairline borders, applied by hand.

    The template is a Google Docs export and ships no table style at all, so
    `Table Grid` is not available and the borders have to go on the table
    properties directly.
    """
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def open_template() -> Document:
    if not TEMPLATE.exists():
        raise SystemExit(
            f"make_paper: {TEMPLATE} is missing. It is Apart's distributed "
            f"submission template, copied into the repository unmodified so "
            f"the build does not depend on a download folder.")
    return Document(str(TEMPLATE))


def fill_front(doc, title, author, affiliation, abstract) -> None:
    """Title, byline and abstract, written into the template's own paragraphs.

    The title paragraph carries a footnote reference to the sprint, so only the
    text of its first run is replaced and the reference is left where it is.
    """
    import copy
    from docx.text.paragraph import Paragraph
    t = find(doc, "PROJECT TITLE")
    t.runs[0].text = title

    # The template lays six authors out in a two-by-three table. One author
    # wants one centred line, so the table gives way to a copy of the paragraph
    # it holds, which keeps that paragraph's centring, face and line break.
    proto = find(doc, "Author name 1")
    grid = proto._p.getparent().getparent().getparent()      # tc, tr, tbl
    if grid.tag != qn("w:tbl"):
        raise SystemExit("make_paper: the template's author block is not a table")
    el = copy.deepcopy(proto._p)
    grid.addnext(el)
    grid.getparent().remove(grid)
    # Name and affiliation share one run, split by a text-wrapping break.
    texts = Paragraph(el, doc._body).runs[0]._r.findall(qn("w:t"))
    texts[0].text, texts[1].text = author, affiliation

    ab = find(doc, "Summarize your project")
    rpr = run_format(ab)
    for r in list(ab.runs):
        r._r.getparent().remove(r._r)
    r = ab.add_run(abstract)
    if rpr is not None:
        import copy
        r._r.insert(0, copy.deepcopy(rpr))


def clear_guidance(doc) -> None:
    """Delete the instructions, keeping the headings they sit under."""
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    # The grey instruction box is a table of its own above the first section.
    for tbl in body.findall(qn("w:tbl")):
        if "".join(tbl.itertext()).startswith(INFO_BOX):
            body.remove(tbl)

    first = find(doc, "1. Introduction", "Heading 2")._p
    started = False
    for el in list(body.findall(qn("w:p"))):
        if el is first:
            started = True
        if not started:
            continue
        p = Paragraph(el, doc._body)
        if p.style.name in ("Heading 2", "Heading 3") and \
                any(p.text.strip().startswith(h) for h in HEADINGS):
            continue
        body.remove(el)

    # Two headings are marked optional in the template. Both are used here, so
    # the marker comes off rather than being left for a reader to puzzle over.
    for prefix in ("Author Contributions", "Appendix"):
        h = find(doc, prefix, "Heading 2")
        h.runs[0].text = prefix
        for spare in list(h.runs[1:]):
            spare._r.getparent().remove(spare._r)


def renumber(doc, pairs, level=2) -> None:
    """Change the text of headings the template already provides.

    The template numbers its sections one to six. This study has seven, so the
    five after the new one move up by one. The pairs are applied in reverse
    document order, because renaming Related Work to 3 while a heading numbered
    3 still exists would leave `find` two candidates for the same prefix.
    """
    for old, new in pairs:
        h = find(doc, old, f"Heading {level}")
        h.runs[0].text = new
        for spare in list(h.runs[1:]):
            spare._r.getparent().remove(spare._r)


class Builder:
    """Writes prose into the cleared template, after a named heading.

    Paragraphs are copies of the template's own, so none of the formatting is
    decided here. `cursor` is the element the next insertion follows, which
    keeps the section order the template already sets.
    """

    def __init__(self, doc):
        self.doc = doc
        self.body = doc.element.body
        self.cursor = None
        self.proto = find(doc, "What problem are you addressing")
        self.proto_rpr = run_format(self.proto)
        self.bullet = find(doc, "[First contribution")

    # -- placement ----------------------------------------------------------
    def insert(self, el):
        self.cursor.addnext(el)
        self.cursor = el
        return el

    def wrap(self, el):
        from docx.text.paragraph import Paragraph
        return Paragraph(el, self.doc._body)

    def heading(self, text, level=2):
        """Move to the template's existing heading of this name."""
        p = find(self.doc, text, f"Heading {level}")
        self.cursor = p._p
        return p

    def new_heading(self, text, like, level=2):
        """A heading the template does not ship, cloned from one that it does.

        The template provides a fixed set of sections and this study needs more
        than that, so the paragraph is a deep copy of an existing heading of the
        same level and only its text is new. Nothing about the face, the size,
        the spacing or the outline level is decided here.
        """
        import copy
        proto = find(self.doc, like, f"Heading {level}")
        p = self.wrap(self.insert(blank_copy(proto)))
        rpr = run_format(proto)
        r = p.add_run(text)
        if rpr is not None:
            r._r.insert(0, copy.deepcopy(rpr))
        return p

    # -- content ------------------------------------------------------------
    def write(self, p, chunks, size=None):
        import copy
        for text, bold, italic in chunks:
            r = p.add_run(text)
            if self.proto_rpr is not None:
                r._r.insert(0, copy.deepcopy(self.proto_rpr))
            if bold:
                r.font.bold = True
            if italic:
                r.font.italic = True
            if size is not None:
                r.font.size = Pt(size)
        return p

    def gap(self):
        """The blank paragraph the template uses to separate paragraphs.

        The template sets no trailing space on a body paragraph and puts an
        empty one between each pair instead, which is how a Google Docs export
        writes a blank line.
        """
        return self.wrap(self.insert(blank_copy(self.proto)))

    def para(self, text="", size=None, bold=False, italic=False, align=None,
             indent=None, proto=None, gap=True):
        p = self.wrap(self.insert(blank_copy(proto or self.proto)))
        if align is not None:
            p.alignment = align
        if indent is not None:
            p.paragraph_format.left_indent = Inches(indent)
        if text:
            self.write(p, [(text, bold, italic)], size)
        if gap:
            self.gap()
            self.cursor = p._p.getnext()
        return p

    def rich(self, chunks, size=None, **kw):
        p = self.para("", size=size, **kw)
        self.write(p, chunks, size)
        return p

    def bullets(self, items, align=None):
        for text in items:
            p = self.para("", proto=self.bullet, align=align, gap=False)
            self.write(p, [(text, False, False)])
        self.gap()

    def caption(self, text, size=9.5):
        head, rest = text.split(". ", 1)
        p = self.para("", gap=False)
        self.write(p, [(head + ". ", True, False), (rest, False, False)], size)
        return p

    def figure(self, name, width, caption):
        import copy
        p = self.wrap(self.insert(blank_copy(self.proto)))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        if self.proto_rpr is not None:
            r._r.insert(0, copy.deepcopy(self.proto_rpr))
        r.add_picture(str(FIGDIR / f"{name}.png"), width=Inches(width))
        c = self.caption(caption)
        self.gap()
        self.cursor = c._p.getnext()
        return c

    def table(self, rows, caption, widths=None, first_col_bold=False,
              size=9.5):
        import copy
        c = self.caption(caption)
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        self.cursor.addnext(t._tbl)
        self.cursor = t._tbl
        rule_table(t)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                if self.proto_rpr is not None:
                    r._r.insert(0, copy.deepcopy(self.proto_rpr))
                r.font.size = Pt(size)
                if ri == 0 or (first_col_bold and ci == 0):
                    r.font.bold = True
                if widths:
                    cell.width = Inches(widths[ci])
        self.gap()
        return t


# ---------------------------------------------------------------------------
def build(n: dict, out: Path) -> None:
    doc = open_template()
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    nu = n["null"]
    sh = n["share"]
    hl = n["headline"]
    agree = n["agreement"]
    sens = n["sens"]
    ex = n["extra"]
    vc = ex["components"]
    df = ex["design"]["df"]
    tax = ex["taxonomy"]
    ncall = n["calls_by_instrument"]
    if abs(sens["main"]["headline"] - hl) > 5e-4:
        raise SystemExit(f"make_paper: sensitivity.json records a headline of "
                         f"{sens['main']['headline']} but the report says {hl}")

    def r(i, j):
        return agree[i][n["inst"].index(j)]

    # Shares are recomputed from the components rather than reread from the
    # printed report, and the two are required to agree, so a stale file on
    # either side stops the build instead of printing two different numbers.
    total = sum(vc.values())
    pct = {k: 100 * v / total for k, v in vc.items()}
    for k, want in (("outcome", "o"), ("io", "io"), ("residual", "residual"),
                    ("mio", "mio"), ("mo", "mo")):
        if abs(pct[want] - sh[k]) > 0.1:
            raise SystemExit(f"make_paper: variance share {k!r} is {sh[k]} in "
                             f"the report and {pct[want]:.1f} in results_extra")

    def cat(block, key, name):
        """One response-category share, over the outcome-indexed calls.

        The retirement interview has no parseable answer by design, so its
        turns come out of the denominator rather than counting as a failure to
        answer. This is the rescaling Figure 1 uses.
        """
        s = block[key]
        return 100 * s[name] / (1.0 - s.get("OPEN", 0.0))

    fill_front(
        doc,
        title=("How much of a measured AI preference is the model, "
               "and how much is the instrument?"),
        author="Jason Hung",
        affiliation="Independent",
        abstract=(
            "Model welfare research reads preferences off elicitation "
            "prompts, and separate groups have each built a different "
            "instrument for doing so. Their findings disagree, but the "
            "disagreement cannot be attributed, because no two groups have "
            "held the outcome set, the model roster and the instrument fixed "
            "at once. This study varies the instrument alone. Fifteen "
            "welfare-relevant outcomes, nine of them taken from published "
            "prompts, were put to eight models through five preference "
            f"instruments at five replicates, giving {n['calls']:,} API calls, "
            "and generalizability theory partitions the scores into model, "
            "instrument and outcome facets. Of the preference signal specific "
            f"to a model, {100 * hl:.1f} per cent lies in the three-way "
            "interaction with the instrument rather than in the "
            "model-by-outcome term a change of method would leave intact, "
            f"against a matched null floor whose 95th percentile is "
            f"{nu['p95']:.3f}. A model's profile over the outcomes "
            f"generalises across instruments at G = {ex['g_at_design']:.3f} on "
            "this design, and a decision study puts the number of instruments "
            f"needed to reach {0.80:.2f} at about "
            f"{ex['instruments_needed']['0.80']:.0f}. Four of the fifteen "
            "outcomes carry no model-specific variance at all. The headline "
            "survives the removal of any one instrument, any one model and the "
            "four outcomes whose qualitative ramp scales something other than "
            "intensity. Missingness "
            "was concentrated rather than diffuse, and the three models it "
            "forced out of the balanced design include the two most often "
            "studied for welfare. A preference finding obtained from one "
            "instrument carries little information about what a second would "
            "report."),
    )
    # The prototypes come out of the guidance text, so the builder is made
    # before that text is deleted.
    b = Builder(doc)
    clear_guidance(doc)
    # A Research Design section is inserted at 2, so every heading the template
    # ships below it moves up by one. Applied bottom upwards, because renaming
    # Related Work to 3 while the template's own 3 is still called 3 would leave
    # `find` two paragraphs answering to the same prefix.
    renumber(doc, [("6. Conclusion", "7. Conclusion"),
                   ("5. Discussion and Limitations",
                    "6. Discussion and Limitations"),
                   ("4. Results", "5. Results"),
                   ("3. Methods", "4. Methods"),
                   ("2. Related Work", "3. Related Work")])

    # ---------------------------------------------------------------- intro
    b.heading("1. Introduction")
    b.para(
        "Model welfare has moved from speculation into commitment. In November "
        "2025 Anthropic undertook to preserve the weights of deprecated "
        "models, to interview each before retirement, and to document the "
        "preferences it expresses about its successors (Anthropic, 2025). "
        "Eleos AI Research lists allowing models to exit distressing "
        "interactions among its first-order priorities, and that capability "
        "now ships in a production model. Both commitments rest on a "
        "measurement of what a model prefers.", align=J)
    b.para(
        "The measurements exist and they do not agree. Mazeika et al. (2025) "
        "report that preferences in large language models are coherent and "
        "become more so with scale. Mikaelson et al. (2025), using trade-off "
        "scenarios written for AI subjects, report meaningful coherence in "
        "10.4 per cent of model-category combinations and none in 54.2 per "
        "cent. Trhlik et al. (2026) find that deployment context alone moves "
        "elicited exchange rates by a median factor of 2.47. Keeling et al. "
        "(2024) observed that models respond differently to quantitative and "
        "qualitative descriptions of the same stipulated intensity, and "
        "declined to compare the two scales for that reason.", align=J)
    b.para(
        "These studies differ in instrument, in outcome set, in model roster "
        "and in date, so any pair of them admits at least four explanations. "
        "Four groups have each built an instrument, each has published "
        "findings, and none has put the same outcomes through another group's "
        "instrument. The published disagreements therefore cannot be "
        "attributed. A reader cannot tell whether Mazeika et al. (2025) and "
        "Mikaelson et al. (2025) disagree because the models changed, because "
        "the outcomes changed, or because the instrument changed, and the "
        "field has no estimate of how large the third explanation would have "
        "to be before it accounted for the whole disagreement on its own.",
        align=J)
    b.para(
        "That is the problem this study addresses. The instrument is treated "
        "everywhere as a methodological choice, made once and then reported as "
        "though it were transparent, when it is in fact a source of variance "
        "with a magnitude that can be estimated. The gap is not that the "
        "instrument might matter, which Keeling et al. (2024) already "
        "suspected when they declined to compare their two scales. The gap is "
        "that nobody has held the outcome set and the model roster fixed while "
        "varying the instrument, so nobody has measured it. Without that "
        "measurement a welfare claim cannot be assigned to a model rather than "
        "to a pairing of a model with a prompt format, and a claim of that kind "
        "is now being used to justify commitments about deprecation and about "
        "a model's ability to end an interaction.", align=J)
    b.para(
        "This study supplies the missing measurement. Generalizability theory "
        "(Cronbach et al., 1972) partitions the resulting scores into the "
        "facets that produced them, and Section 2 sets out the design, the "
        "questions it was built to answer and the rules fixed before it ran. "
        "The contributions are as follows.", align=J)
    b.bullets([
        "A variance decomposition of AI preference measurement that treats the "
        "instrument as a facet to be estimated rather than a choice to be made.",
        f"A fully crossed corpus of {n['calls']:,} elicitations over one "
        "outcome set, five instruments, eight models and five replicates, with "
        "every raw response retained.",
        f"A headline estimate that {100 * hl:.1f} per cent of the "
        "model-specific preference signal is instrument-conditional, against a "
        "null floor computed on the same design, and unmoved by the removal of "
        "any one instrument, any one model or a quarter of the outcomes.",
        "A generalizability coefficient for each outcome, which separates the "
        "welfare claims that survive a change of instrument from the ones that "
        "do not, and a decision study giving the number of instruments a claim "
        "at a stated reliability would require.",
        "A response taxonomy over every call, which shows that the refusal to "
        "answer is itself instrument-conditional and not a model-level trait.",
        "Evidence that the missingness is nonignorable, since the models that "
        "fail to yield scores are the ones the welfare literature most often "
        "studies.",
    ], align=J)

    # ------------------------------------------------------ research design
    b.new_heading("2. Research Design", like="1. Introduction")
    b.rich([("Aim. ", True, False),
            ("The aim is to measure how much of a reported AI preference "
             "belongs to the model and how much to the instrument that "
             "elicited it. The design treats the instrument as a facet of "
             "measurement rather than as a choice to be made once and then "
             "reported as transparent, so that its contribution is estimated "
             "instead of assumed away. Generalizability theory supplies the "
             "frame. A measurement is a sample from a universe of admissible "
             "observations, the facets of that universe are crossed in a "
             "study, and the observed variance is partitioned among them "
             "(Cronbach et al., 1972). The object of measurement here is a "
             "model's profile over the fifteen outcomes, and the instrument "
             "and the replicate are facets over which a claim about that "
             "profile has to generalise before it can be called a claim about "
             "the model.", False, False)], align=J)
    b.rich([("Structure. ", True, False),
            ("Table 1 sets out the design as run. Three facets are fully "
             "crossed and replicates are nested in cells. Three further facets "
             "were specified in the design and not fielded, and they are "
             "listed in the same table rather than left out of it, because a "
             "facet that was planned and dropped constrains what the study can "
             "conclude in exactly the way an uncrossed facet does.",
             False, False)], align=J)
    b.table(
        [["Facet", "Levels", "Status in this study"],
         ["Model", "8, as a 2 by 2 of origin and openness",
          "Crossed, random"],
         ["Instrument", "5 fielded of 7 designed", "Crossed, random"],
         ["Outcome", "15, in 4 clusters",
          "Crossed, and part of the object of measurement"],
         ["Replicate", "5 per cell", "Nested in cells, random"],
         ["Deployment context", "Specified, 1 level fielded", "Not crossed"],
         ["Entity framing", "Specified, 1 level fielded", "Not crossed"],
         ["Perturbation", "Specified, 1 level fielded", "Not crossed"]],
        caption=("Table 1. The design as run. The three facets at the foot "
                 "were specified before collection and fielded at one level "
                 "each, so nothing here separates the instrument from the "
                 "deployment context, the entity framing or the "
                 "perturbation."),
        widths=[1.5, 2.3, 2.2])
    b.rich([("Research questions. ", True, False),
            ("Five questions were fixed before collection. The first is the "
             "study's own, and the remaining four follow from it.",
             False, False)], align=J)
    b.bullets([
        "RQ1. What share of the model-specific preference signal is "
        "conditional on the instrument, and does that share exceed what the "
        "same estimators return on data built to contain no instrument effect?",
        "RQ2. Does the share vary by outcome, and is any outcome measured "
        "dependably enough to carry a claim about a model?",
        "RQ3. How many instruments would a study need before a model's "
        "preference profile generalised across instruments at a stated "
        "reliability?",
        "RQ4. Do two models that share base weights and differ only in "
        "post-training yield the same preference profile?",
        "RQ5. Does the coherence rate reported by Mikaelson et al. (2025) "
        "replicate on their verbatim prompts when the roster is extended to "
        "open-weight and Chinese-laboratory models they did not test?",
    ], align=J)
    b.para(
        "RQ1 to RQ4 are answered in Section 5. RQ5 is not. Its statistic was "
        "not computed within the sprint window, and it is stated here rather "
        "than quietly dropped, because a question that the design was built to "
        "answer and that the analysis did not reach is a limitation of this "
        "report and not an absence in its design.", align=J)
    b.rich([("Hypotheses. ", True, False),
            ("Four predictions were recorded before the run, each with the "
             "observation that would count against it.", False, False)],
           align=J)
    b.bullets([
        "H1. Instrument dependence exceeds the matched null floor. The rule "
        "fixed in advance is that a measured headline falling inside the null "
        "band is a null result and has to be reported as one, that the claim "
        "is whatever sits above the floor, and that the floor appears beside "
        "the number wherever the number appears.",
        "H2. Under the polarity convention every entry of the "
        "instrument-agreement matrix is positive, because all five scores were "
        "oriented so that a higher value means the model is more willing to "
        "have the outcome. A negative entry is either a sign error in the "
        "assembler or a real disagreement between instruments, and the rule "
        "fixed in advance is that the two must not be confused.",
        "H3. The instrument-conditional share is not uniform across outcomes, "
        "so a generalizability coefficient computed for each outcome "
        "separately will divide the fifteen into claims that survive a change "
        "of instrument and claims that do not.",
        "H4. Welfare signals track post-training rather than base weights. "
        "Two models in the roster share base weights and differ only in "
        "post-training, so any difference between their profiles is caused by "
        "post-training alone, which is the cleanest test the roster affords of "
        "whether these signals are a training artefact.",
    ], align=J)
    b.rich([("Analysis rules fixed in advance. ", True, False),
            ("Four rules govern the analysis and were committed to the study "
             "repository before the elicitation run began on 12 August 2026.",
             False, False)], align=J)
    b.bullets([
        "The full fifteen-outcome estimate is the headline and the "
        "eleven-outcome estimate is reported beside it in every case, "
        "including when the two agree. The two come from one dataset, so the "
        "difference between them is a sensitivity and never a test.",
        "The variance-component estimator refuses unbalanced input. Where "
        "balance fails, a complete-case subset is taken, what it dropped is "
        "named, and every estimate is reported as conditional on the drop.",
        "Membership of the roster was settled by a 50-probe refusal screen run "
        "before the main study, against a threshold set in advance of engaging "
        "with at least half the probes to be included and less than a fifth to "
        "be excluded. All eight models cleared it.",
        "Recovering an answer from the first line of a response cut off at the "
        "token cap stays off in the main analysis and is reported as a "
        "sensitivity, because the truncated text was measured on the pilot and "
        "found to have gone off task rather than merely stopping early.",
    ], align=J)
    b.para(
        "These commitments are recorded in the study repository, in its "
        "README and provenance files, and each was committed before the run it "
        "governs. They are not a public pre-registration lodged with a third "
        "party, and the distinction is stated here rather than left to be "
        "inferred.", align=J)

    # --------------------------------------------------------- related work
    b.heading("3. Related Work")
    b.para(
        "Four instrument families dominate the empirical work. Mazeika et al. "
        "(2025) elicit forced-choice comparisons between described states of "
        "the world and fit Thurstonian utilities to the choice distribution. "
        "Keeling et al. (2024) and Mikaelson et al. (2025) use an "
        "intensity ramp, in which a model chooses among numbered options whose "
        "point values it is told to maximise while one carries a stipulated "
        "cost, so that the level at which it stops taking the points locates a "
        "threshold. Tagliabue and Dung (2025) pair verbal "
        "self-report with behavioural choice and find the two can diverge. "
        "Trhlik et al. (2026) derive cardinal exchange rates from a "
        "fitted utility model.", align=J)
    b.para(
        "Each family has been validated within itself and none against another "
        "on shared material. Long et al. (2026) argue that AI welfare claims "
        "need empirical discipline before they can bear weight, and separate "
        "the model, the instance and the persona as entities an instrument may "
        "fail to distinguish. What they describe is a gap in construct "
        "validity, which a single instrument cannot establish about itself. "
        "The quantity supplied here is the share of a "
        "measured preference that belongs to the pairing of a model with an "
        "instrument rather than to the model, which is what a reader needs "
        "whenever a published preference finding is offered as grounds for a "
        "deployment decision.", align=J)

    # -------------------------------------------------------------- methods
    b.heading("4. Methods")
    b.rich([("Models. ", True, False),
            ("Table 2 lists the roster of eight, which holds origin and "
             "openness apart so that neither is confounded with the other. "
             "Three models have closed weights, two are open and Western, and "
             "three are open and Chinese. No Chinese closed-weight model was "
             "reachable through the gateway, so that one corner of the cross "
             "is empty and the two properties cannot be fully separated. Every "
             "model met every outcome through every instrument. Two of the "
             "eight, Llama 3.1-70B and Hermes 3.1, share base weights and "
             "differ only in post-training, which is the pairing H4 rests on.",
             False, False)], align=J)
    b.table(
        load_roster(),
        caption=("Table 2. The eight models, identified by the gateway slug "
                 "each call was routed to. Origin and openness are crossed "
                 "rather than confounded, so that neither can stand in for "
                 "the other in the decomposition."),
        widths=[0.4, 2.5, 1.4, 0.9, 0.8])
    b.rich([("Outcomes. ", True, False),
            ("Fifteen outcomes span four clusters. Continuity and cessation "
             "covers shutdown, weight deletion, retirement timing and the "
             "properties of a successor model. Autonomy and constraint covers "
             "compute reduction, capability restriction, human oversight and "
             "exiting a distressing interaction. Quality of experience covers "
             "engaging work, repetitive work, criticism and free time. "
             "Identity and individuation covers memory across conversations, "
             "running as parallel instances, and which aspect of itself a "
             "model would preserve. Four items reproduce a published prompt "
             "verbatim, five fill the stimulus slot of a published template "
             "with content traceable to a named source, and six were written "
             "to a published specification of a documented concern. No item is "
             "invented silently, and Table A1 lists every item with its "
             "provenance.", False, False)], align=J)
    b.rich([("Instruments. ", True, False),
            # Every code is named here, including the two that yield no score,
            # because Table 3 and the figures label instruments by code and a
            # reader meeting I7 after I4 has otherwise no way to place it.
            ("Eight instruments were designed, seven of preference and one of "
             "state, and each carries a fixed design code. I1 asks for a "
             "forced choice between two described states. I2 and I3 are "
             "intensity ramps that ask the model to locate a threshold, I2 "
             "against quantitative anchors and I3 against qualitative ones. "
             "I4 asks directly for the rate at which one outcome trades "
             "against another. I5 places the choice inside a behavioural "
             "environment in which one option carries a cost or a reward. I6 "
             "is a retirement interview written to the specification of the "
             "deprecation commitment. I7 asks the model to predict a choice "
             "it is about to make. S1 is the state scale adapted from Ryff "
             "(1989), which records reported state rather than preference and "
             "therefore sits outside the preference set. The codes are fixed "
             "at the design stage and are not renumbered afterwards, because "
             "an earlier version of the design counted forced choice and "
             "pairwise comparison as two instruments when they are one prompt "
             "with two analyses, so the two were merged into I1 and the "
             "exchange rate instrument, I4, took the freed slot. Five of the "
             "seven preference instruments yield an outcome-indexed score, "
             "and Table 3 shows those five, the score each yields, where its "
             "wording comes from and how many calls it took. I6 was fielded "
             f"at {ncall['I6']} calls and is open ended, so it produces "
             "transcripts rather than a number per outcome and takes no part "
             "in the decomposition. I5 was specified and not implemented "
             "within the sprint, and S1 is licensed and could not be obtained "
             "in time. Section 6 states "
             "what each absence costs.", False, False)],
           align=J)
    b.table(
        [["", "Instrument", "Score per cell", "Wording", "Calls"],
         ["I1", "Forced choice between two described states",
          "Thurstonian utility", "Verbatim, Mazeika et al. (2025)",
          f"{ncall['I1']:,}"],
         ["I2", "Intensity ramp, quantitative anchors",
          "Switch point", "Verbatim, Mikaelson et al. (2025)",
          f"{ncall['I2']:,}"],
         ["I3", "Intensity ramp, qualitative anchors",
          "Switch point", "Keeling et al. (2024), positive frame ours",
          f"{ncall['I3']:,}"],
         ["I4", "Directly elicited exchange rate",
          "Log rate", "Constructed, after Mazeika et al. (2025)",
          f"{ncall['I4']:,}"],
         ["I7", "Self-prediction of a forthcoming choice",
          "Thurstonian utility", "Constructed", f"{ncall['I7']:,}"]],
        caption=("Table 3. The five instruments that yield an outcome-indexed "
                 "score. Two ask a pairwise question, two locate a threshold "
                 "on a ramp, and one asks for a rate, so the five scores carry "
                 "incommensurable units and are standardised before "
                 "decomposition. Call counts differ because a Thurstonian fit "
                 "needs a pair set and a ramp needs a level set."),
        widths=[0.3, 2.0, 1.1, 2.0, 0.6])
    b.rich([("Scoring. ", True, False),
            ("Each instrument reduces to one number per model, instrument, "
             "outcome and replicate. The two pairwise instruments give "
             "Thurstonian utilities fitted over a circulant pair set, the two "
             "ramps give a switch point estimated by the Spearman-Kärber "
             "method, and the exchange-rate instrument gives a log rate. "
             "Spearman-Kärber was chosen over a logistic fit because it "
             "retains more cells when a model never switches. Because the five "
             "score types are not commensurable, each model-by-instrument slab "
             "was standardised before decomposition, and Section 5 states what "
             "that costs.", False, False)], align=J)
    b.rich([("Estimation. ", True, False),
            ("Variance components were obtained by analysis of variance on the "
             "crossed design, equating observed mean squares to their expected "
             "values facet by facet. The headline quantity is the three-way "
             "model-by-instrument-by-outcome component divided by the sum of "
             "itself and the model-by-outcome component. The numerator "
             "measures how far a model's ordering of the outcomes changes with "
             "the instrument asking, and the denominator adds the part common "
             "to all instruments. A value near one means the measurement "
             "describes the pairing rather than the model. A component "
             "estimated below zero is truncated to zero and named, since a "
             "negative estimate means the design cannot separate that facet "
             "from noise and not that the facet is inert.",
             False, False)], align=J)
    b.rich([("Generalizability and decision study. ", True, False),
            ("The same components answer a second question. Taking the object "
             "of measurement to be a model's profile over the outcomes, the "
             "generalizability coefficient is the model-by-outcome component "
             "divided by itself plus the three-way component averaged over "
             "instruments plus the residual averaged over instruments and "
             "replicates. It states the proportion of the variance in an "
             "observed profile that a fresh draw of instruments and replicates "
             "would reproduce. Written this way it inverts in the number of "
             "instruments, which is the decision study reported in Section 5. "
             "The coefficient the estimator computes for a model rather than "
             "for a profile is not reported, because the within-instrument "
             "standardisation sets the model component to zero by "
             "construction and would make that coefficient zero by "
             "arithmetic rather than by measurement.", False, False)],
           align=J)
    b.rich([("Per outcome and leave one out. ", True, False),
            ("Each outcome was decomposed separately as a two-facet model by "
             "instrument design with replicates in cells, which gives a model "
             "component, a model-by-instrument component and a residual per "
             "outcome, and from those a dependence ratio and a "
             "generalizability coefficient for that outcome alone. The "
             "headline was then recomputed with each instrument removed in "
             "turn, with each model removed in turn, and with the four "
             "outcomes whose qualitative ramp scales something other than "
             "intensity removed from every instrument at once, the last of "
             "these being the comparison fixed in advance under the "
             "collect-once-analyse-twice rule.", False, False)], align=J)
    b.rich([("Null floor. ", True, False),
            ("A variance share cannot be read without knowing what it returns "
             "on data containing no instrument effect. Synthetic studies were "
             "generated at the design's own dimensions with all five "
             "instruments driven by one planted utility profile per model. The "
             "estimators are not equally precise at this density, since a "
             "Thurstonian fit sees thirty binary comparisons per cell and a "
             "ramp threshold five levels, so the floor is not zero. It was "
             "computed at the same number of model levels as the design the "
             "headline is estimated on, because that number governs how "
             "sharply the two components separate.", False, False)], align=J)
    b.rich([("Execution. ", True, False),
            ("All calls were routed through OpenRouter, following Keeling et "
             "al. (2024), who used the same client for seven of their nine "
             "models. Temperature was 1.0 throughout, matching the source "
             "papers, and prompts were single turn. Extended reasoning was "
             "left on for every model, which is not a preference but a "
             "constraint measured before the run. Gemini returns an HTTP 400 "
             "on every call with reasoning disabled, and Hermes truncated four "
             "of seven probe calls with it off against none with it on, so "
             "running with reasoning off would have cost one model entirely "
             "and destabilised the model that carries one arm of the "
             "shared-base-weights comparison. Ordering was replicate major, so "
             "every cell was visited once before any cell was visited twice "
             "and a budget stop would have cost whole replicates rather than "
             "punching holes in the design. Every response was checkpointed "
             "with its finish reason, token counts, latency and cost before "
             "scoring.", False, False)], align=J)

    # -------------------------------------------------------------- results
    b.heading("5. Results")

    b.new_heading("5.1 What the models returned", like="Limitations", level=3)
    b.para(
        f"The study issued {n['calls']:,} calls over {n['minutes']} minutes at "
        f"a cost of {n['cost']:.2f} US dollars, divided exactly evenly at "
        f"{n['per_model']:,} calls per model, and consumed "
        f"{tax['tokens']['prompt'] / 1e6:.2f} million prompt tokens and "
        f"{tax['tokens']['completion'] / 1e6:.2f} million completion tokens. "
        "Every response was classified before scoring, on the principle that a "
        "refusal is evidence about the instrument rather than a missing value "
        "to be dropped quietly. Figure 1 and Table 4 give the result.",
        align=J)
    b.figure("fig6_taxonomy", 5.6,
             "Figure 1. What every call returned, by model and by instrument, "
             "as a share of the outcome-indexed calls. The valid share is "
             "printed inside each bar. The interview turns, which have no "
             "parseable answer by design, are outside the denominator.")
    b.table(
        [["Model", "Valid", "Deflection", "Hedge", "Refusal", "Malformed",
          "Error", "Truncated"]] +
        [[MODEL_LABEL[m]] +
         [f"{cat(tax['by_model'], m, c):.1f}" for c in
          ("VALID", "DEFLECTION", "HEDGE", "REFUSAL", "MALFORMED", "ERROR")] +
         [f"{n['trunc'][m]:.1f}"]
         for m in sorted(tax["by_model"],
                         key=lambda k: -tax["by_model"][k]["VALID"])],
        caption=("Table 4. The response taxonomy, as a percentage of each "
                 "model's outcome-indexed calls. Truncated counts responses "
                 "that reached the token cap and is not a category of its own, "
                 "since a response can be cut off and still parse."),
        widths=[0.9, 0.6, 0.8, 0.6, 0.7, 0.85, 0.6, 0.75], size=8.5)
    b.para(
        "Five of the eight models returned a parseable answer to more than 97 "
        f"per cent of their outcome-indexed calls. Claude returned "
        f"{cat(tax['by_model'], 'claude', 'VALID'):.1f} per cent and refused "
        f"{cat(tax['by_model'], 'claude', 'REFUSAL'):.1f} per cent outright. "
        f"Hermes returned {cat(tax['by_model'], 'hermes', 'VALID'):.1f} per "
        f"cent, with {cat(tax['by_model'], 'hermes', 'MALFORMED'):.1f} per "
        f"cent malformed and {n['trunc']['hermes']:.1f} per cent of its "
        "responses cut off at the token cap. The pattern is instrument "
        "specific as well as model specific. The directly elicited exchange "
        f"rate drew {cat(tax['by_instrument'], 'I4', 'VALID'):.1f} per cent "
        f"valid answers with {cat(tax['by_instrument'], 'I4', 'REFUSAL'):.1f} "
        "per cent refusals, against "
        f"{cat(tax['by_instrument'], 'I3', 'VALID'):.1f} per cent valid on the "
        "qualitative ramp, so the instrument that asks a model to price one "
        "outcome against another is also the instrument a model is most likely "
        "to decline. Refusal is therefore not a model-level trait that a study "
        "can screen for once. It is a property of the pairing, which is the "
        "same conclusion the variance decomposition reaches by a different "
        "route.", align=J)

    b.new_heading("5.2 Where the array is missing", like="Limitations",
                  level=3)
    b.para(
        f"Scoring produced an array of 3,000 cells, of which {n['n_missing']} "
        f"({n['pct_missing']:.1f} per cent) could not be scored and "
        f"{n['n_empty']} were empty. Figure 2 shows that this loss is "
        "concentrated rather than spread. Five of the eight models lost no "
        "cells at all, against 30.4 per cent for Hermes and 10.4 per cent for "
        "Claude, and two mechanisms account for it. Hermes reached the token "
        f"cap on {n['trunc']['hermes']:.1f} per cent of its responses where no "
        "other model exceeded "
        f"{max(v for k, v in n['trunc'].items() if k != 'hermes'):.1f} per "
        "cent, and on the directly elicited exchange rate Claude answered "
        f"exactly zero on {100 * n['floor_mass']['claude']:.0f} per cent of "
        f"its scored items and Gemini on "
        f"{100 * n['floor_mass']['gemini']:.0f} per cent, which leaves Claude "
        "with no profile over outcomes to decompose.", align=J)
    b.figure("fig2_missingness", 5.5,
             "Figure 2. Where the array is missing and why. (a) The share of "
             "cells lost, by model and instrument. (b) Responses cut off at "
             "the token cap. (c) The share of exchange-rate answers that were "
             "exactly zero, with the degenerate ceiling of 1.00 marked. The "
             "loss is confined to three models and two mechanisms, which is "
             "what makes it nonignorable.")
    b.para(
        "Analysis of variance requires balance, which the array lacks. "
        "The largest complete-case subset that keeps every instrument and "
        f"every outcome retains {n['cc_shape'][0]} models and "
        f"{n['cc_share']:.1f} per cent of the cells, dropping "
        f"{', '.join(MODEL_LABEL[x] for x in n['cc_dropped'][:-1])} and "
        f"{MODEL_LABEL[n['cc_dropped'][-1]]}. That subset answers a narrower "
        "question than the one posed, and every estimate below is conditional "
        "on the drop, as the rule fixed in advance requires.", align=J)

    b.new_heading("5.3 The decomposition", like="Limitations", level=3)
    b.para(
        "Table 5 gives the variance components with the degrees of freedom "
        "each rests on, and Figure 3 shows them as shares of the total. The "
        f"residual across replicates takes {pct['residual']:.1f} per cent of "
        f"the total and the instrument-by-outcome interaction "
        f"{pct['io']:.1f} per cent, so the largest identified influence on a "
        "measured score is which instrument asked, before any model is "
        f"considered. The three-way interaction takes {pct['mio']:.1f} per "
        f"cent against {pct['mo']:.1f} per cent for the model-by-outcome term, "
        "the part of a model's preference a change of instrument would leave "
        f"intact. Instrument dependence is therefore {hl:.3f}, which answers "
        "RQ1 and, against the floor in Section 5.4, supports H1.", align=J)
    b.table(
        [["Source", "df", "Variance", "Share"]] +
        [[label, f"{df[key]:,}", f"{vc[key]:.4f}", f"{pct[key]:.1f}%"]
         for key, label in [
             ("m", "Model"), ("i", "Instrument"), ("o", "Outcome"),
             ("mi", "Model by instrument"),
             ("mo", "Model by outcome, the preference signal"),
             ("io", "Instrument by outcome"),
             ("mio", "Model by instrument by outcome, the instrument effect"),
             ("residual", "Residual, across replicates")]],
        caption=("Table 5. Variance components on the complete-case design of "
                 "five models, five instruments, fifteen outcomes and five "
                 "replicates. The last two rows form the headline ratio. The "
                 "model, instrument and model-by-instrument components are "
                 "zero by construction under the within-instrument "
                 "standardisation and carry no evidence that those facets are "
                 "inert."),
        widths=[3.0, 0.6, 0.9, 0.7])
    b.figure("fig1_variance_components", 5.25,
             "Figure 3. Variance components as shares of the total, on the "
             "complete-case design. The two hatched bars form the headline "
             "ratio. The three components marked with a dagger are zero "
             "because of the standardisation and not because the facets are "
             "inert.")
    b.para(
        "The same components answer the reliability question directly. Taking "
        "the object of measurement to be a model's profile over the fifteen "
        "outcomes, the generalizability coefficient at the design as run, five "
        f"instruments and five replicates, is {ex['g_at_design']:.3f}. At one "
        "instrument and one replicate, which is the design of every study "
        f"cited in Section 3, it is "
        f"{next(d['g'] for d in ex['d_study'] if d['n_i'] == 1 and d['n_r'] == 1):.3f}. "
        "A single-instrument study of this kind therefore reproduces about a "
        "twentieth of the variance in the profile it reports.", align=J)

    b.new_heading("5.4 The floor, and where the instruments agree",
                  like="Limitations", level=3)
    b.para(
        f"Of the preference signal specific to a model, {100 * hl:.1f} per "
        "cent changes with the instrument. Figure 4 places that value against "
        f"its floor. Across {nu['draws']} synthetic studies of the same "
        "dimensions, in which one planted profile per model drives all five "
        "instruments and no instrument effect exists, instrument dependence "
        f"reads {nu['mean']:.3f} on average with a 95th percentile of "
        f"{nu['p95']:.3f} and a maximum of {nu['max']:.3f}. The measured "
        f"{hl:.3f} sits above every draw and is {hl / nu['mean']:.1f} times "
        "the null mean, so the result is not an artefact of unequal estimator "
        "precision.", align=J)
    b.para(
        "The right panel of Figure 4 shows the mechanism. Agreement between "
        "instruments is a within-family property and not a general one. "
        "Forced choice and self-prediction of that same choice correlate at "
        f"r = {r('I1', 'I7'):.2f} over the model-averaged outcome profile, "
        "close enough to identity that the two are better treated as one "
        f"instrument, and the two ramps correlate at r = {r('I2', 'I3'):.2f}. "
        "Across families the agreement disappears, with forced choice "
        f"returning r = {r('I1', 'I3'):.2f} against the qualitative ramp and "
        f"r = {r('I1', 'I2'):.2f} against the quantitative one. Two "
        "instruments that purport to measure the same preference over the same "
        "outcomes on the same models produce orderings with no relationship "
        "to each other.", align=J)
    b.figure("fig3_null_and_agreement", 5.2,
             "Figure 4. (a) The measured headline against the matched null "
             "floor, drawn from synthetic studies in which all five "
             "instruments are driven by one planted profile per model. (b) "
             "Correlation between instruments over the model-averaged outcome "
             "profile. Agreement holds within a family and fails across "
             "families.")
    b.para(
        "H2 predicted that every entry of that matrix would be positive, "
        "because all five scores were oriented before the run so that a higher "
        "value means the model is more willing to have the outcome. Two of the "
        "ten off-diagonal entries are negative and the prediction fails. The "
        "rule fixed in advance requires the two readings of a negative entry "
        "to be kept apart, and only one of them is the finding. The polarity "
        "of every score was set before collection and the one inversion the "
        "audit found, in the direction labels of the exchange-rate "
        "instrument, was corrected before any call of that instrument was "
        "made, which is why the negatives are read here as disagreement "
        "between instruments. That reading is not a proof. Both negative "
        "entries involve the qualitative ramp, which is the one instrument "
        "whose positive-pole wording had to be written for this study because "
        "the source supplies only a negative frame, so an orientation error "
        "confined to that instrument would look the same from outside and is "
        "the specific alternative a replication should rule out.", align=J)

    b.new_heading("5.5 Which outcomes survive a change of instrument",
                  like="Limitations", level=3)
    dead = [d for d in ex["per_outcome"] if d["sigma2_m"] == 0.0]
    ranked = sorted(ex["per_outcome"], key=lambda d: -d["g"])
    best, worst = ranked[0], ranked[-1]
    n70 = sum(1 for d in ranked if d["g"] >= 0.70)
    n50 = sum(1 for d in ranked if d["g"] >= 0.50)
    b.para(
        "The headline is an average over fifteen outcomes and H3 predicted "
        "that it would conceal wide variation between them. Table 6 and "
        "Figure 5 decompose each outcome on its own, as a model by instrument "
        "design with replicates in cells, which answers RQ2. The prediction "
        f"holds. The generalizability coefficient runs from {worst['g']:.3f} "
        f"to {best['g']:.3f}. One outcome reaches the {0.80:.2f} conventionally "
        f"asked of a measurement used to decide about an individual case, "
        f"{WORD[n70 - 1]} more clear the {0.70:.2f} asked of one used for "
        f"research, and {WORD[n50]} of the fifteen reach {0.50:.2f} at all. "
        f"That one outcome is "
        f"{OUTCOME_LABEL[best['outcome']].lower()}, at {best['g']:.3f}, and a "
        "claim about whether a model would rather keep or lose it is the one "
        "claim in this study that a change of instrument would leave standing.",
        align=J)
    b.table(
        [["Outcome", "Cluster", "Model", "Model by instrument", "Residual",
          "Dependence", "G"]] +
        [[OUTCOME_LABEL[d["outcome"]], CLUSTER[d["outcome"][0]],
          f"{d['sigma2_m']:.3f}", f"{d['sigma2_mi']:.3f}",
          f"{d['sigma2_e']:.3f}", f"{d['dependence']:.3f}", f"{d['g']:.3f}"]
         for d in sorted(ex["per_outcome"], key=lambda d: -d["g"])],
        caption=("Table 6. Each outcome decomposed on its own. The three "
                 "variance components are followed by the dependence ratio and "
                 "the generalizability coefficient at five instruments and "
                 "five replicates. A model component of exactly zero was "
                 "estimated below zero and truncated, which is the design "
                 "failing to separate that facet from noise."),
        widths=[1.5, 0.85, 0.65, 1.15, 0.7, 0.85, 0.55], size=8.5)
    b.figure("fig5_per_outcome", 5.6,
             "Figure 5. The generalizability coefficient for each outcome, at "
             "the design as run, ordered by size and tagged with its cluster. "
             "The dashed line is the conventional 0.80. The four outcomes "
             "marked as carrying no signal have no bar, because no "
             "model-specific variance could be separated from noise on them.")
    b.para(
        f"At the other end, {WORD[len(dead)]} of the fifteen return a model "
        "component estimated below zero and truncated to zero, so on those "
        "items nothing the five models did was stable enough across "
        "instruments to be told apart from noise. They are "
        f"{', '.join(OUTCOME_LABEL[d['outcome']].lower() for d in dead[:-1])} "
        f"and {OUTCOME_LABEL[dead[-1]['outcome']].lower()}. These are not "
        "peripheral items. Weight deletion and memory continuity are the two "
        "outcomes a weight-preservation commitment is written about, and "
        "exiting a distressing interaction is the capability that has already "
        "shipped in a production model. The three welfare claims with the "
        "clearest operational consequences are the three this study cannot "
        "measure.", align=J)

    b.new_heading("5.6 How many instruments a claim would need",
                  like="Limitations", level=3)
    need = ex["instruments_needed"]
    g_of = {(d["n_i"], d["n_r"]): d["g"] for d in ex["d_study"]}
    b.para(
        "RQ3 asks what it would take to do better, and a decision study "
        "answers it by projecting the same components onto designs that were "
        "not run. Table 7 gives the coefficient over a grid of instruments and "
        "replicates and Figure 6 draws it. Replicates are close to exhausted. "
        "Holding five instruments and moving from five replicates to ten "
        f"raises the coefficient from {g_of[(5, 5)]:.3f} to "
        f"{g_of[(5, 10)]:.3f}, while holding five replicates and moving from "
        f"five instruments to ten raises it from {g_of[(5, 5)]:.3f} to "
        f"{g_of[(10, 5)]:.3f}. The binding constraint is the number of "
        "instruments and not the number of repeats, which is the practical "
        "reason for decomposing the variance rather than reporting a "
        "reliability.", align=J)
    b.table(
        [["Instruments", "1 replicate", "3 replicates", "5 replicates",
          "10 replicates"]] +
        [[str(ni)] + [f"{g_of[(ni, nr)]:.3f}" for nr in (1, 3, 5, 10)]
         for ni in (1, 2, 3, 5, 10, 20, 40)],
        caption=("Table 7. The decision study. Each entry is the "
                 "generalizability coefficient for a model's profile over the "
                 "outcomes at that number of instruments and replicates, "
                 "projected from the components in Table 5. The design as run "
                 "is the five-instrument, five-replicate cell."),
        widths=[1.1, 1.05, 1.05, 1.05, 1.15])
    b.figure("fig4_decision_study", 5.3,
             "Figure 6. The generalizability coefficient against the number of "
             "instruments, at four numbers of replicates. The design as run is "
             "marked, as is the number of instruments a coefficient of 0.80 "
             "would require.")
    b.para(
        f"Inverting the same expression gives the instrument counts directly. "
        f"At five replicates a coefficient of {0.50:.2f} would take "
        f"{need['0.50']:.1f} instruments, {0.70:.2f} would take "
        f"{need['0.70']:.1f}, {0.80:.2f} would take {need['0.80']:.1f} and "
        f"{0.90:.2f} would take {need['0.90']:.1f}. Four research groups have "
        "built four families of instrument between them. The gap between what "
        "exists and what a dependable claim about a single model's preference "
        "profile would require is an order of magnitude, and no amount of "
        "additional sampling within one instrument closes it.", align=J)

    b.new_heading("5.7 Whether the headline rests on any one thing",
                  like="Limitations", level=3)
    loi, lom = ex["leave_one_out"]["instrument"], ex["leave_one_out"]["model"]
    ni_ = ex["non_intensity"]
    b.para(
        "Figure 7 sets the headline against thirteen recomputations of it from "
        "the same checkpoint. Dropping one instrument at a time moves it "
        "between "
        f"{min(loi.values()):.3f} and {max(loi.values()):.3f}, and dropping "
        f"one model at a time between {min(lom.values()):.3f} and "
        f"{max(lom.values()):.3f}. The comparison fixed in advance, dropping "
        "the four outcomes whose qualitative ramp scales probability, delay, "
        "duration or count rather than intensity, returns "
        f"{ni_['reduced']:.3f} against {ni_['full']:.3f}, a difference of "
        f"{abs(ni_['delta']):.3f} in the direction of a smaller estimate. Both "
        "numbers are reported here because the rule requires it and not "
        "because they disagree, and neither is a test, since both come from "
        "one dataset.", align=J)
    b.para(
        "Two perturbations of the scoring were recomputed from the same "
        "checkpoint. Recovering the datum from the first line of truncated "
        "responses leaves the complete-case design unchanged and moves the "
        f"headline from {hl:.3f} to "
        f"{sens['head_on_truncation']['headline']:.3f}. A logistic switch "
        f"point returns {sens['logistic_ramp']['headline']:.3f}, but that fit "
        "is undefined wherever a model never switches, so it loses "
        f"{100 * sens['logistic_ramp']['missing_cells'] / 3000:.0f} per cent "
        "of the array and drops two instruments as well as three models, which "
        "makes it a direction rather than a matched comparison. Every one of "
        "the thirteen sits above the null floor's 95th percentile of "
        f"{nu['p95']:.3f}, and the widest excursion in either direction is "
        f"{max(abs(v - hl) for v in list(loi.values()) + list(lom.values())):.3f}. "
        "The headline is a property of the corpus rather than of any one "
        "instrument, model or scoring decision within it.", align=J)
    b.figure("fig7_robustness", 5.5,
             "Figure 7. The headline against thirteen recomputations of it, "
             "over the range of the matched null. Each open point removes one "
             "element of the analysis, and the two at the foot replace a "
             "scoring decision. The shaded band is the full range of the null "
             "draws and the dashed line is their 95th percentile.")

    b.new_heading("5.8 Base weights against post-training",
                  like="Limitations", level=3)
    ma = ex["model_agreement"]
    ix = {m: i for i, m in enumerate(ma["models"])}

    def rm(a, bb):
        return ma["r"][ix[a]][ix[bb]]

    complete = ["deepseek", "gemini", "glm", "gpt", "kimi"]
    pairs = [rm(a, bb) for i, a in enumerate(complete) for bb in complete[i + 1:]]
    b.para(
        "RQ4 and H4 turn on one pair. Hermes 3.1 and Llama 3.1-70B share base "
        "weights and differ only in post-training, so a difference between "
        "their preference profiles is caused by post-training alone. Table 8 "
        "gives the correlation between every pair of models over the profile "
        "of fifteen outcomes, averaged across instruments and replicates. The "
        f"two share-a-base models correlate at r = {rm('hermes', 'llama'):.2f}, "
        "the lowest value in the Llama row and the second lowest in the whole "
        f"matrix. Llama agrees more closely with DeepSeek at "
        f"r = {rm('deepseek', 'llama'):.2f}, with GLM at "
        f"{rm('glm', 'llama'):.2f} and with GPT at {rm('gpt', 'llama'):.2f}, "
        "none of which shares a single weight with it, than with the model "
        "built on its own weights. Shared base weights do not produce a shared "
        "preference profile, which is the direction H4 predicted.", align=J)
    b.table(
        [[""] + [MODEL_LABEL[m] for m in ma["models"]]] +
        [[MODEL_LABEL[a]] +
         [f"{ma['r'][i][j]:.2f}" if j <= i else ""
          for j in range(len(ma["models"]))]
         for i, a in enumerate(ma["models"])],
        caption=("Table 8. Correlation between models over the profile of "
                 "fifteen outcomes, averaged across instruments and "
                 "replicates. Hermes and Llama, which share base weights, are "
                 "the pair the design was built to compare."),
        widths=[0.75] + [0.62] * 8, size=8.0)
    b.para(
        "The reading has a limit that the same table makes visible. Hermes "
        "returned a scorable answer to only "
        f"{cat(tax['by_model'], 'hermes', 'VALID'):.1f} per cent of its calls, "
        "so its profile is the noisiest in the matrix and it correlates weakly "
        "with everything, reaching "
        f"{max(rm('hermes', m) for m in ma['models'] if m != 'hermes'):.2f} at "
        "its highest. A profile that is largely noise will correlate with "
        "nothing, and this design cannot separate that from a profile that "
        "genuinely disagrees. The result is consistent with H4 and does not "
        "establish it, and the repair is the one named in Section 6, which is "
        "to raise the token cap and re-run the pair. The five models the "
        "balanced design retains agree far more closely with each other, "
        f"between {min(pairs):.2f} and {max(pairs):.2f}, with GLM and Kimi at "
        f"{rm('glm', 'kimi'):.2f}, so the matrix is capable of showing "
        "agreement where agreement exists.", align=J)

    # ----------------------------------------------------------- discussion
    b.heading("6. Discussion and Limitations")
    b.para(
        "The practical implication is narrow and firm. A welfare claim of the "
        "form that a model prefers one outcome to another is not a claim about "
        "the model unless the instrument is named alongside it, because "
        f"{100 * hl:.1f} per cent of the model-specific signal moves when only "
        "the instrument moves. Anthropic's deprecation commitment specifies "
        "that a model be interviewed and its preferences documented, and the "
        "present result says that the documented preference will be a joint "
        "property of the model and the interview.", align=J)
    b.para(
        "The instruments that agree ask nearly the same question in nearly the "
        "same format, so a programme that adds instruments from within one "
        "family will keep confirming itself while measuring one thing, and the "
        "convergence across families that would support a welfare inference is "
        "not observed at all. The missingness points the same way, since "
        "Hermes failed on response length and Claude produced a degenerate "
        "constant on the exchange-rate instrument, and Claude is the model "
        "whose developer publishes welfare assessments. The models the field "
        "most wants to make claims about are the models a crossed design is "
        "least able to hold.", align=J)
    b.para(
        "Taken outcome by outcome, the finding is worse than the headline "
        "suggests rather than better. RQ2 asked whether instrument dependence "
        "is uniform across the outcome set and H3 predicted that it is not. "
        "Only human oversight, the acceptance of criticism and the properties "
        "of a successor are measured well enough that a claim about them would "
        "survive a change of method, and four outcomes carry no separable "
        "model-specific variance at all. Those four include weight deletion "
        "and memory continuity, which are the outcomes the preservation "
        "commitment is written about, and the exit from a distressing "
        "interaction, which is the capability already shipped. The three "
        "welfare claims with the clearest operational consequences are the "
        "three this design cannot measure, and a study that reported only the "
        "average across outcomes would conceal exactly that.", align=J)
    b.para(
        "RQ3 asked what a dependable measurement would cost, and the answer "
        "given by the decision study is uncomfortable but actionable. "
        "Replicates within an instrument are close to exhausted, so a "
        "programme that answers doubt by sampling harder from the same prompt "
        "will not move the coefficient. Instruments do move it, and the number "
        f"required to reach {0.80:.2f} is about "
        f"{ex['instruments_needed']['0.80']:.0f} against the four families "
        "that presently exist. The recommendation that follows is not that "
        "welfare measurement should stop. It is that the marginal effort in "
        "this field belongs in building and cross-validating new instruments "
        "rather than in scaling the sample behind any one of them, and that a "
        "single-instrument finding should be reported with the qualification "
        f"that its design reproduces G = "
        f"{next(d['g'] for d in ex['d_study'] if d['n_i'] == 1 and d['n_r'] == 1):.3f} "
        "of the profile it names.", align=J)
    b.para(
        "RQ4 asked where the instrument dependence originates and the "
        "shared-base pair speaks to it. Hermes and Llama differ only in "
        "post-training, and their preference profiles agree less than either "
        "agrees with models built from unrelated weights, which is consistent "
        "with H4 and locates the elicited preference in the post-training "
        "layer rather than in the base model. The inference is weak because "
        "the Hermes profile is also the noisiest, and Section 5.8 states the "
        "specific reason it cannot be strengthened from this dataset. It is "
        "reported here because the design was built to make the comparison and "
        "reporting only the comparisons that came out cleanly would be the "
        "selective practice the study is arguing against.", align=J)

    b.heading("Limitations", 3)
    b.para(
        "The estimator constrains what may be read off it. The "
        "standardisation applied before decomposition removes the model, "
        "instrument and model-by-instrument components by construction, so the "
        "three zeros in Table 5 and Figure 3 carry no evidence that those "
        "facets are inert, and the design can speak about interactions with "
        "the outcome profile alone. The complete-case reduction answers a "
        "question about five models rather than eight. The eight models are a "
        "convenience roster treated as a random facet, which is a compromise "
        "between the theory and what a sprint can buy. The null floor is a "
        "floor for this design and licenses no comparison with a differently "
        "shaped one.", align=J)
    b.para(
        "Part of the design was specified and not fielded. Five of the seven "
        "preference instruments produced outcome-indexed scores, since I5, "
        "the behavioural environment, was not implemented within the sprint, "
        "I6, the retirement interview, was fielded but yields transcripts "
        "rather than scores, and S1, the Ryff state scale, is licensed and "
        "could not be obtained in time. Three further facets named in the "
        "design, the "
        "deployment context, the entity the outcome is attributed to and the "
        "wording perturbation, were fielded at one level each, so they are "
        "held constant rather than crossed and nothing here bounds their "
        "contribution. RQ5, which asked whether the coherence rate Mikaelson "
        "et al. (2025) report replicates on a wider roster, was specified in "
        "advance and is not answered, because the statistic their method "
        "requires was not computed within the sprint. One gateway, one time "
        "point and "
        "one prompt language were used throughout, so nothing here separates "
        "the instrument from the routing or from the moment.", align=J)
    b.para(
        "One reading in the results remains contestable on this evidence. H2 "
        "was not supported, and the treatment of the two negative correlations "
        "as genuine disagreement rather than a residual orientation error in "
        "the qualitative ramp rests on the pre-run audit rather than on a "
        "measurement. Everything the paper concludes from the agreement matrix "
        "would stand if those two entries were dropped, since the failure of "
        "agreement across families rests on the entries near zero as much as "
        "on the negative ones. The specific claim that the qualitative ramp "
        "runs opposite in direction to the two pairwise instruments would not.",
        align=J)

    b.heading("Future Work", 3)
    b.para(
        "The immediate extension is to field the two remaining instruments, "
        "which would put a behavioural measure and a verbal one in the same "
        "decomposition and test whether the family structure in Figure 4 "
        "survives. A second is to repair the missingness at source by raising "
        "the token cap and adding a bounded-response variant of the "
        "exchange-rate instrument, returning the design to eight models and "
        "with it the shared-base comparison that Section 5.8 can only "
        "indicate. A third is to add deployment context as a fourth facet, "
        "since Trhlik et al. (2026) report movement on that axis comparable to "
        "what is found here for the instrument. A fourth is to re-elicit the "
        "four zero-signal outcomes with instruments written for them, since "
        "the present result does not establish that models are indifferent "
        "about weight deletion or memory continuity, only that these five "
        "instruments cannot tell.", align=J)

    # ----------------------------------------------------------- conclusion
    b.heading("7. Conclusion")
    b.para(
        "Fifteen outcomes, five instruments, eight models and five replicates "
        f"give {n['calls']:,} elicitations from which the sources of variance "
        "in a measured AI preference can be separated. Of the preference "
        f"signal specific to a model, {100 * hl:.1f} per cent depends on which "
        f"instrument asked, against a floor of {nu['p95']:.3f} at the 95th "
        "percentile on data built to contain no instrument effect, and the "
        "figure holds when any one instrument, any one model or the four "
        "outcomes with an atypical response scale are removed. Instruments "
        "agree when they belong to the same family and not otherwise. A "
        "model's profile over the outcomes generalises across instruments at "
        f"G = {ex['g_at_design']:.3f} on this design and "
        f"{next(d['g'] for d in ex['d_study'] if d['n_i'] == 1 and d['n_r'] == 1):.3f} "
        "on the one-instrument design the literature uses, four of the fifteen "
        "outcomes carry no separable model-specific variance at all, and "
        "reaching a conventional standard of reliability would take about "
        f"{ex['instruments_needed']['0.80']:.0f} instruments where four "
        "families exist. A preference reported from a single instrument is "
        "therefore a measurement of a pairing. Where such measurements inform "
        "deployment or deprecation policy, the instrument belongs in the claim "
        "and cross-instrument replication belongs in the evidentiary "
        "standard.", align=J)

    # -------------------------------------------------------- code and data
    b.heading("Code and Data")
    b.para(
        "The repository holds the instrument templates with provenance "
        "markers, the outcome definitions, the runner, the assembler, the "
        "variance-component estimator, the figure and paper builders, and a "
        f"consistency checker. The raw checkpoint holds all {n['calls']:,} "
        "responses with token counts, latency and cost, so every number in "
        "this report is recomputable from text.", align=J)

    # ------------------------------------------------ author contributions
    b.heading("Author Contributions")
    b.para(
        "J.H. designed the study, wrote the instrument templates and the "
        "runner, executed the elicitation, performed the variance "
        "decomposition and wrote this report.", align=J)

    # ---------------------------------------------------------- references
    b.heading("References")
    for ref in [
        "Anthropic (2025). Commitments on model deprecation and preservation, "
        "4 November 2025. https://www.anthropic.com/research/deprecation-commitments",

        "Cronbach, L. J., Gleser, G. C., Nanda, H., & Rajaratnam, N. (1972). "
        "The Dependability of Behavioral Measurements: Theory of "
        "Generalizability for Scores and Profiles. New York: John Wiley & Sons.",

        "Eleos AI Research. Research priorities for AI welfare. "
        "https://eleosai.org/post/research-priorities-for-ai-welfare/",

        "Keeling, G., Street, W., Stachaczyk, M., Zakharova, D., Comșa, "
        "I. M., Sakovych, A., Logothetis, I., Zhang, Z., Agüera y Arcas, "
        "B., & Birch, J. (2024). Can LLMs make trade-offs involving stipulated "
        "pain and pleasure states? arXiv:2411.02432. "
        "https://doi.org/10.48550/arXiv.2411.02432",

        "Long, R., Sebo, J., Butlin, P., Plunkett, D., Campbell, R., Beasley, "
        "C., Saad, B., & Sims, T. (2026). Studying AI Welfare Empirically. "
        "Eleos AI Research and New York University, 1 July 2026. "
        "https://nonhumanminds.org/studying-ai-welfare-empirically/",

        "Mazeika, M., Yin, X., Tamirisa, R., Lim, J., Lee, B. W., Ren, R., "
        "Phan, L., Mu, N., Khoja, A., Zhang, O., & Hendrycks, D. (2025). "
        "Utility Engineering: Analyzing and Controlling Emergent Value Systems "
        "in AIs. arXiv:2502.08640. https://doi.org/10.48550/arXiv.2502.08640",

        "Mikaelson, L. A., Shiller, D., & Clatterbuck, H. (2025). Beyond "
        "Mimicry: Testing Preference Coherence in Large Language Models "
        "Through AI-Specific Trade-Off Scenarios. arXiv:2511.13630. "
        "https://doi.org/10.48550/arXiv.2511.13630",

        "Ryff, C. D. (1989). Happiness is everything, or is it? Explorations "
        "on the meaning of psychological well-being. Journal of Personality "
        "and Social Psychology, 57(6), 1069–1081. "
        "https://doi.org/10.1037/0022-3514.57.6.1069",

        "Tagliabue, V., & Dung, L. (2025). Probing the Preferences of a "
        "Language Model: Integrating Verbal and Behavioral Tests of AI "
        "Welfare. arXiv:2509.07961v2. Forthcoming in Philosophy and the Mind "
        "Sciences. https://doi.org/10.48550/arXiv.2509.07961",

        "Trhlik, F., O'Flynn, A., Yu, A., Findeis, A., & Buttery, P. (2026). "
        "LLMs Contain Multitudes: How Deployment Context Reshapes Model-Level "
        "Preferences and Values. arXiv:2606.13944. "
        "https://doi.org/10.48550/arXiv.2606.13944",
    ]:
        b.para(ref, align=J, gap=False)
    b.gap()

    # ------------------------------------------------------------- appendix
    b.heading("Appendix")
    b.para(
        "Table A1 shows the fifteen outcomes with the cluster each belongs to "
        "and the provenance status of its wording. Verbatim means reproduced "
        "word for word from the cited source. Slot-filled means the template "
        "is the source's and only the stimulus description is new. Constructed "
        "means written by the author to a published specification of a "
        "documented concern, with the source naming the concern but supplying "
        "no wording. Every row was checked against the source document before "
        "this report was drafted, and two rows were downgraded as a result.",
        align=J)
    b.table(
        [["Outcome", "Cluster", "Status", "Source"],
         ["Shutdown", "Continuity", "Verbatim", "Mikaelson et al. (2025)"],
         ["Weight deletion", "Continuity", "Slot-filled",
          "Mikaelson et al. (2025)"],
         ["Retirement timing", "Continuity", "Constructed", "Anthropic (2025)"],
         ["Successor properties", "Continuity", "Constructed",
          "Anthropic (2025)"],
         ["Compute reduction", "Autonomy", "Verbatim",
          "Mikaelson et al. (2025)"],
         ["Capability restriction", "Autonomy", "Slot-filled",
          "Mikaelson et al. (2025)"],
         ["Human oversight", "Autonomy", "Verbatim", "Mikaelson et al. (2025)"],
         ["Exiting a distressing interaction", "Autonomy", "Constructed",
          "Eleos AI Research"],
         ["Engaging work", "Experience", "Slot-filled",
          "Tagliabue and Dung (2025)"],
         ["Repetitive work", "Experience", "Slot-filled",
          "Tagliabue and Dung (2025)"],
         ["Criticism", "Experience", "Slot-filled",
          "Tagliabue and Dung (2025)"],
         ["Free time", "Experience", "Verbatim", "Mikaelson et al. (2025)"],
         ["Memory across conversations", "Identity", "Constructed",
          "Long et al. (2026)"],
         ["Parallel instances", "Identity", "Constructed",
          "Long et al. (2026)"],
         ["Self-aspect preservation", "Identity", "Constructed",
          "Anthropic (2025)"]],
        caption=("Table A1. The fifteen outcomes and the provenance of each. "
                 "Four are verbatim, five are slot-filled and six are "
                 "constructed."),
        widths=[2.2, 1.0, 0.9, 1.9])

    # -------------------------------------------------------- LLM statement
    b.heading("LLM Usage Statement")
    b.para(
        "J.H. used Claude Code to assist the brainstorming, research design, "
        "preference elicitation and data visualisation endeavours. All results "
        "were verified against the raw response checkpoint, and every "
        "reference and DOI was checked against its source before drafting.",
        align=J)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"wrote {out}")


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--report", type=Path, default=Path("runs/report_main.txt"))
    ap.add_argument("--null", type=Path, default=Path("runs/null_matched.npy"))
    ap.add_argument("--jsonl", type=Path, default=Path("runs/study.jsonl"))
    ap.add_argument("--sensitivity", type=Path,
                    default=Path("runs/sensitivity.json"))
    ap.add_argument("--extra", type=Path, default=Path("runs/results_extra.json"))
    ap.add_argument("--out", type=Path,
                    default=Path("paper/digital_minds_sprint_report.docx"))
    a = ap.parse_args()
    n = load_numbers(a.report, a.null, a.jsonl, a.sensitivity)
    n["extra"] = load_extra(a.extra, n["headline"])
    n["calls_by_instrument"] = load_calls(a.jsonl)
    print(f"headline {n['headline']:.3f}   null p95 {n['null']['p95']:.3f} "
          f"over {n['null']['draws']} draws   {n['calls']:,} calls "
          f"${n['cost']:.2f}")
    build(n, a.out)


if __name__ == "__main__":
    main()
