"""Read the built .docx back and check it against the rules it was written to.

This exists because the rules are easy to state and easy to violate one word at
a time. It reads the saved file rather than the source strings, so it also
catches anything the template contributed.

    python3 check_paper.py paper/digital_minds_sprint_report.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

# Punctuation that must not appear in running prose. The em and en dash are
# ruled out outright; the colon is allowed only where it belongs to a source's
# real title or to a URL, both of which are checked for separately below.
BANNED = {"—": "em dash", "–": "en dash", "―": "horizontal bar"}

# Phrasing that reads as machine-written. Each was chosen because it appears in
# generated text far more often than in published academic prose.
TICS = [
    r"\bdelve\b", r"\bdelves\b", r"\bit is worth noting\b", r"\bcrucially\b",
    r"\bimportantly\b", r"\bnotably\b", r"\bmoreover\b", r"\bfurthermore\b",
    r"\badditionally\b", r"\bin conclusion\b", r"\bin summary\b",
    r"\boverall,", r"\bleverage\b", r"\brobustly\b", r"\bseamless\b",
    r"\bunderscore\b", r"\bunderscores\b", r"\bhighlights the importance\b",
    r"\bplays a (?:crucial|vital|key) role\b", r"\brich tapestry\b",
    r"\bnavigat(?:e|ing) the\b", r"\bit is important to note\b",
    r"\bnot only .{0,40} but also\b", r"\bcornerstone\b", r"\brealm\b",
    r"\bparadigm shift\b", r"\bgame.chang", r"\bdive into\b",
    r"\bshed light on\b", r"\bat the end of the day\b", r"\bthat being said\b",
]


def paragraphs(doc):
    """Every paragraph in the document, however deeply nested.

    Walking `doc.paragraphs` and `doc.tables` misses anything inside a table
    within a table, and Apart's template puts the byline and the whole abstract
    three levels down, so that route left the abstract unchecked.
    """
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for el in doc.element.body.iter(f"{ns}p"):
        yield "".join(t.text or "" for t in el.iter(f"{ns}t"))


def main() -> int:
    path = Path(sys.argv[1] if len(sys.argv) > 1
                else "paper/digital_minds_sprint_report.docx")
    doc = Document(str(path))
    texts = [t for t in paragraphs(doc) if t.strip()]
    body = "\n".join(texts)
    fails = []

    for ch, name in BANNED.items():
        for t in texts:
            if ch in t:
                # A page range in a reference is typography, not phrasing.
                if name == "en dash" and re.search(r"\d–\d", t):
                    continue
                fails.append(f"{name} in: {t[:90]}")

    # Colons are permitted in a URL and inside a reference entry, where they
    # belong to a source's real title or to a publisher's city, and nowhere
    # else. Reference entries are the paragraphs after the References heading.
    try:
        ref_start = texts.index("References")
        ref_end = next(i for i, t in enumerate(texts)
                       if i > ref_start and t == "Appendix")
    except (ValueError, StopIteration):
        ref_start, ref_end = len(texts), len(texts)
        fails.append("could not locate the References section")
    for k, t in enumerate(texts):
        if ref_start < k < ref_end:
            continue
        for m in re.finditer(r":", t):
            i = m.start()
            if re.search(r"https?:|arXiv:|doi:", t[max(0, i - 8):i + 4], re.I):
                continue
            fails.append(f"colon in: {t[max(0, i - 60):i + 40]}")

    for pat in TICS:
        for t in texts:
            if re.search(pat, t, re.I):
                fails.append(f"phrasing {pat!r} in: {t[:80]}")

    # Every exhibit must be introduced by name in the running text.
    captions = {}
    for t in texts:
        m = re.match(r"((?:Figure|Table) A?\d+)\.", t.strip())
        if m:
            captions[m.group(1)] = t
    for name, caption in sorted(captions.items()):
        refs = [t for t in texts
                if re.search(rf"\b{name}\b", t) and t != caption]
        if not refs:
            fails.append(f"{name} has a caption but is never introduced in text")

    # Numbers that must agree with the run, wherever they appear.
    if "11,528" not in body:
        fails.append("the run size 11,528 does not appear")
    # 11,528 calls, 1,441 per model, a 3,000-cell array of which 2,411 survive,
    # the per-instrument call counts, and the residual degrees of freedom.
    known = ("11,528", "1,441", "3,000", "2,411", "2,400", "1,200", "1,500")
    for stray in re.findall(r"\b\d{1,3},\d{3}\b", body):
        if stray not in known:
            fails.append(f"unexplained four-figure number {stray}")

    # The template's own scaffolding must be gone.
    for marker in ["How to use this template", "[First contribution",
                   "PROJECT TITLE", "[Reference 1]", "e.g., \"A.B. led"]:
        if marker in body:
            fails.append(f"template guidance left in place: {marker!r}")

    words = len(re.findall(r"\S+", body))
    print(f"{path}  {len(texts)} paragraphs, {words} words, "
          f"{len(captions)} exhibits {sorted(captions)}")
    if fails:
        print(f"\n{len(fails)} problem(s):")
        for f in dict.fromkeys(fails):
            print("  -", f)
        return 1
    print("clean")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
